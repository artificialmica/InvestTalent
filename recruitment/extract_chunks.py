"""
Noun Chunk Extraction Script - PRODUCTION VERSION
Extracts all noun chunks from CVs for classifier training
Supports: .txt, .pdf, .docx

Usage:
    python extract_chunks.py --input_dir /path/to/cvs --output chunks_dataset.csv

Requirements:
    pip install spacy pandas pdfplumber python-docx
    python -m spacy download en_core_web_sm
"""

import sys
import os
from pathlib import Path
import argparse
import re

# ----------------------------
# Check and import dependencies
# ----------------------------
try:
    import pandas as pd
except ImportError:
    print("[!] pandas not installed. Run: pip install pandas")
    sys.exit(1)

try:
    import spacy
except ImportError:
    print("[!] spacy not installed. Run: pip install spacy")
    sys.exit(1)

# PDF support
try:
    import pdfplumber
    PDF_SUPPORT = "pdfplumber"
except ImportError:
    try:
        import PyPDF2
        PDF_SUPPORT = "PyPDF2"
    except ImportError:
        PDF_SUPPORT = None
        print("[!] Warning: No PDF library found. Install with: pip install pdfplumber")

# DOCX support
try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("[!] Warning: python-docx not found. Install with: pip install python-docx")

# ----------------------------
# Load spaCy model
# ----------------------------
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    print("Downloading spaCy model...")
    os.system("python -m spacy download en_core_web_sm")
    try:
        nlp = spacy.load("en_core_web_sm")
    except OSError:
        print("[!] Failed to load spaCy model. Run: python -m spacy download en_core_web_sm")
        sys.exit(1)


# ============================================================
#  Text Extraction - Supports TXT, PDF, DOCX
# ============================================================
def extract_text_from_pdf(filepath):
    """Extract text from PDF file."""
    if PDF_SUPPORT == "pdfplumber":
        try:
            text = ""
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            print(f"[!] pdfplumber error on {filepath}: {e}")
            return ""
    
    elif PDF_SUPPORT == "PyPDF2":
        try:
            text = ""
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            return text
        except Exception as e:
            print(f"[!] PyPDF2 error on {filepath}: {e}")
            return ""
    
    else:
        print(f"[!] No PDF library available - skipping {filepath}")
        print("    Install with: pip install pdfplumber")
        return ""


def extract_text_from_docx(filepath):
    """Extract text from DOCX file."""
    if not DOCX_SUPPORT:
        print(f"[!] python-docx not installed - skipping {filepath}")
        print("    Install with: pip install python-docx")
        return ""
    
    try:
        doc = DocxDocument(filepath)
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"[!] Error reading DOCX {filepath}: {e}")
        return ""


def extract_text_from_file(filepath):
    """Extract text from .txt, .pdf, or .docx."""
    filepath = str(filepath)
    
    try:
        if filepath.lower().endswith(".txt"):
            return Path(filepath).read_text(encoding="utf-8", errors="ignore")

        elif filepath.lower().endswith(".pdf"):
            return extract_text_from_pdf(filepath)

        elif filepath.lower().endswith(".docx"):
            return extract_text_from_docx(filepath)

        else:
            # Try reading as text
            return Path(filepath).read_text(encoding="utf-8", errors="ignore")

    except Exception as e:
        print(f"[!] Error reading {filepath}: {e}")
        return ""


# ============================================================
# Section Detection (Europass-friendly)
# ============================================================
EUROPASS_SECTIONS = {
    "work experience": "experience",
    "professional experience": "experience",
    "employment history": "experience",

    "education": "education",
    "education and training": "education",

    "skills": "skills",
    "digital skills": "skills",
    "personal skills": "skills",
    "communication skills": "skills",
    "interpersonal skills": "skills",
    "technical skills": "skills",
    "key competencies": "skills",

    "certifications": "certifications",
    "certificates": "certifications",
    
    "languages": "languages",
    "profile": "profile",
    "summary": "profile",
}


def detect_section(text_before):
    """Detect which section the noun chunk belongs to."""
    ctx = text_before[-300:].lower()

    # Europass section keywords
    for key, value in EUROPASS_SECTIONS.items():
        if key in ctx:
            return value

    # Generic fallback
    if any(x in ctx for x in ["skills:", "technical skills"]):
        return "skills"
    if any(x in ctx for x in ["experience", "employment"]):
        return "experience"
    if any(x in ctx for x in ["education", "academic"]):
        return "education"

    return "unknown"


# ============================================================
# Chunk Cleaning Heuristics
# ============================================================
STOPWORDS = set([
    "and", "or", "the", "a", "an", "in", "for", "on", "of", "to", "at",
    "by", "with", "about", "into", "research", "journal", "conference",
    "study", "paper", "chapter", "section", "i", "my", "we", "our", "they"
])

FORBIDDEN_KEYWORDS = [
    "journal", "study", "conference", "paper", "chapter",
    "investigator", "coordinator"
]


def is_valid_chunk(text):
    """Filters obviously invalid noun chunks BEFORE training."""
    text = text.strip()
    words = text.lower().split()

    # Too short or too long
    if len(words) < 1 or len(words) > 5:
        return False

    # Single character or very short
    if len(text) < 2:
        return False

    # Single word that is a stopword
    if len(words) == 1 and words[0] in STOPWORDS:
        return False

    # All words are stopwords
    if all(w in STOPWORDS for w in words):
        return False

    # Contains forbidden patterns
    if any(k in text.lower() for k in FORBIDDEN_KEYWORDS):
        return False

    # No alphabetic characters
    if not re.search(r"[a-zA-Z]", text):
        return False

    # Just numbers or dates (e.g., "2022", "01/2023")
    if re.match(r"^[\d\s\-/\.,:]+$", text):
        return False

    return True


# ============================================================
# Chunk Extraction
# ============================================================
def extract_chunks_from_cv(filepath, cv_id):
    """Extract all noun chunks from a single CV."""
    text = extract_text_from_file(filepath)
    if not text:
        print(f"[!] No text extracted from {filepath}")
        return []

    print(f"Processing: {Path(filepath).name} ({len(text)} chars)")

    # Process with spaCy (limit to 50k chars for safety)
    doc = nlp(text[:50000])

    chunks_output = []
    seen = set()

    for chunk in doc.noun_chunks:
        chunk_text = chunk.text.strip()

        # Deduplicate (case-insensitive)
        if chunk_text.lower() in seen:
            continue
        seen.add(chunk_text.lower())

        # Validate structure
        if not is_valid_chunk(chunk_text):
            continue

        # Position and section
        pos = chunk.start_char
        section = detect_section(text[:pos])

        # POS pattern normalization
        pos_tags = [tok.pos_ for tok in chunk]
        pos_pattern = "-".join(pos_tags)

        chunks_output.append({
            "chunk_id": f"{cv_id}_{len(chunks_output)}",
            "text": chunk_text,
            "cv_id": cv_id,
            "source_file": Path(filepath).name,
            "section": section,
            "word_count": len(chunk_text.split()),
            "char_count": len(chunk_text),
            "pos_pattern": pos_pattern,
            "has_digit": any(c.isdigit() for c in chunk_text),
            "has_special": any(c in chunk_text for c in ["'", '"', '-', '–']),
            "label": None  # To be filled during labeling
        })

    print(f"  -> {len(chunks_output)} valid chunks extracted\n")
    return chunks_output


# ============================================================
# Main
# ============================================================
def main():
    parser = argparse.ArgumentParser(description="Extract noun chunks from CVs for ML training")
    parser.add_argument("--input_dir", default=".", help="Folder containing CVs")
    parser.add_argument("--output", default="chunks_dataset.csv", help="Output CSV file")
    parser.add_argument("--max_cvs", type=int, default=20, help="Maximum CVs to process")

    args = parser.parse_args()

    input_dir = Path(args.input_dir)
    
    if not input_dir.exists():
        print(f"[!] Directory not found: {input_dir}")
        sys.exit(1)

    # Find all CV files
    cv_files = []
    for ext in ("*.txt", "*.pdf", "*.docx", "*.PDF", "*.DOCX"):
        cv_files.extend(list(input_dir.glob(ext)))

    if not cv_files:
        print(f"[!] No CV files found in {input_dir}")
        print("    Looking for: .txt, .pdf, .docx files")
        sys.exit(1)

    cv_files = cv_files[:args.max_cvs]
    
    print("=" * 60)
    print("CHUNK EXTRACTION")
    print("=" * 60)
    print(f"Found {len(cv_files)} CV files in {input_dir}\n")
    
    # Check dependencies
    print("Dependencies:")
    print(f"  PDF support: {PDF_SUPPORT if PDF_SUPPORT else 'NOT AVAILABLE'}")
    print(f"  DOCX support: {'Available' if DOCX_SUPPORT else 'NOT AVAILABLE'}")
    print()

    all_chunks = []

    for idx, cv in enumerate(cv_files, start=1):
        chunks = extract_chunks_from_cv(str(cv), f"cv{idx:03d}")
        all_chunks.extend(chunks)

    if not all_chunks:
        print("[!] No chunks extracted. Check your CV files and dependencies.")
        sys.exit(1)

    # Save to CSV
    df = pd.DataFrame(all_chunks)
    df.to_csv(args.output, index=False)

    # Summary
    print("=" * 60)
    print("EXTRACTION COMPLETE")
    print("=" * 60)
    print(f"Total chunks extracted: {len(df)}")
    print(f"From {len(cv_files)} CVs")
    print(f"Saved to: {args.output}")
    
    # Section breakdown
    print("\nChunks by section:")
    for section in df['section'].value_counts().items():
        print(f"  {section[0]}: {section[1]}")
    
    print("\n" + "=" * 60)
    print("NEXT STEPS")
    print("=" * 60)
    print("1. Open the CSV file")
    print("2. Add labels in the 'label' column:")
    print("   - 1 = This is a skill")
    print("   - 0 = This is NOT a skill")
    print("3. Label at least 200 chunks")
    print(f"4. Run: python train_classifier.py --input {args.output}")


if __name__ == "__main__":
    main()